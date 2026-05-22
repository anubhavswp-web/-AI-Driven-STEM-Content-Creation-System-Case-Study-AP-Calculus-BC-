import os
import zipfile
import shutil
import tempfile
import re
import subprocess
from tkinter import Tk, filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

def insert_para_after(paragraph):
    """Low-level XML injection to robustly insert a paragraph exactly after another."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    return new_para

def get_clean_id(text):
    """Extracts alphanumeric identity for universal matching across all Calculus units."""
    text = text.lower()
    identity = ""
    if 'prob' in text: identity += 'prob'
    if 'pract' in text: identity += 'pract'
    identity += "".join(re.findall(r'\d+', text))
    return identity

def place_images_automatically():
    root = Tk()
    root.withdraw() 
    root.attributes("-topmost", True)

    # 1. Select the Source Document
    messagebox.showinfo("Step 1", "Select your AP Calculus Word Document (.docx)")
    docx_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if not docx_path: return

    # 2. Select the Image ZIP
    messagebox.showinfo("Step 2", "Select the ZIP file of images")
    zip_path = filedialog.askopenfilename(filetypes=[("Zip files", "*.zip")])
    if not zip_path: return

    # 3. Generate the Automatic Output Path
    doc_dir = os.path.dirname(docx_path)
    doc_name = os.path.basename(docx_path)
    name_part, extension = os.path.splitext(doc_name)
    default_output_name = f"{name_part} with images{extension}"
    output_path = os.path.join(doc_dir, default_output_name)

    temp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        image_files = [f for f in os.listdir(temp_dir) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]
        
        doc = Document(docx_path)
        paragraphs = list(doc.paragraphs)
        placed_count = 0

        print(f"Analyzing: {doc_name}...")

        for paragraph in paragraphs:
            para_text = paragraph.text.strip()
            if not para_text or len(para_text) < 2: continue

            clean_para = get_clean_id(para_text)
            if not clean_para: continue

            for img_name in image_files:
                clean_img = get_clean_id(img_name)

                if clean_img == clean_para:
                    img_path = os.path.join(temp_dir, img_name)
                    
                    img_para = insert_para_after(paragraph)
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_para.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    
                    cap_para = insert_para_after(img_para)
                    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    c_run = cap_para.add_run(f"Figure: Visual representation for {para_text}")
                    c_run.italic = True
                    c_run.font.name = 'Arial'
                    
                    placed_count += 1
                    break 

        doc.save(output_path)
        
        # 4. Success and Open Folder Logic
        success_msg = f"Task Complete!\nPlaced {placed_count} images.\nSaved to: {default_output_name}"
        messagebox.showinfo("Success", success_msg)
        
        # Opens file explorer and selects the file
        subprocess.run(['explorer', '/select,', os.path.normpath(output_path)])

    except Exception as e:
        messagebox.showerror("Error", f"Technical Error: {str(e)}")
    finally:
        shutil.rmtree(temp_dir)
        root.destroy()

if __name__ == "__main__":
    place_images_automatically()