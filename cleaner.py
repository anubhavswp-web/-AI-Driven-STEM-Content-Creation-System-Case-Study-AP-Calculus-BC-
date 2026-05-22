import os
import tkinter as tk
from tkinter import filedialog, simpledialog
from docx import Document

def clean_documents():
    # 1. Setup UI
    root = tk.Tk()
    root.withdraw()

    print("--- 🧹 Calculus Document Symbol Cleaner ---")
    
    # 2. Select Folder
    target_folder = filedialog.askdirectory(title="Select the Unit Folder to Clean")
    if not target_folder:
        print("Selection cancelled.")
        return

    # 3. ASK USER FOR SYMBOLS (The new feature)
    # Example input: ###, [10pt], \[[10pt]]
    user_input = simpledialog.askstring("Input", "Enter symbols to remove (separated by commas):",
                                        initialvalue="###, [10pt], \\[10pt]")
    
    if not user_input:
        print("No symbols entered. Exiting.")
        return

    # Convert the string into a list and clean up spaces
    unwanted_symbols = [s.strip() for s in user_input.split(",")]
    print(f"Targeting these symbols: {unwanted_symbols}")
    
    files_cleaned = 0

    # 4. Process Files
    for root_dir, dirs, files in os.walk(target_folder):
        for filename in files:
            if filename.endswith(".docx") and not filename.startswith("~$"):
                file_path = os.path.join(root_dir, filename)
                
                try:
                    doc = Document(file_path)
                    is_modified = False

                    # Clean Paragraphs
                    for para in doc.paragraphs:
                        for symbol in unwanted_symbols:
                            if symbol in para.text:
                                para.text = para.text.replace(symbol, "")
                                is_modified = True

                    # Clean Tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    for symbol in unwanted_symbols:
                                        if symbol in para.text:
                                            para.text = para.text.replace(symbol, "")
                                            is_modified = True

                    if is_modified:
                        doc.save(file_path)
                        print(f"✨ Cleaned: {filename}")
                        files_cleaned += 1
                        
                except Exception as e:
                    print(f"❌ Error in {filename}: {e}")

    print(f"\n✅ SUCCESS: {files_cleaned} files were cleaned.")
    input("Press Enter to close...")

if __name__ == "__main__":
    clean_documents()