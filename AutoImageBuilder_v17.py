"""
AutoImageBuilder v17
====================

BASE: v9 STABLE architecture (proven to work)
  - One AI call for whole document
  - OMML math equation reader (exact functions, no guessing)
  - 450 DPI Ultra-HD rendering
  - Temp file → insert → cleanup
  - JSON sanitiser
  - 4-6 line descriptions

NEW in v17 (added on top):
  - Background thread  ─ GUI never freezes
  - Progress bar       ─ see exactly where it is
  - Blank image check  ─ retries if image renders empty
  - Code sanitiser     ─ fixes plt.title/xlabel/ylabel, bad escapes, sqrt(-x)
  - Deep visual prompt ─ accurate geometry, MVT diagrams, labeled dimensions,
                         exact functions, secant + tangent lines
  - Folder mode fix    ─ processes ALL files (not just first)
  - Model fallbacks    ─ tries 4 models before giving up

Requirements (pip install these):
  google-genai  python-docx  customtkinter  matplotlib  numpy  Pillow

Usage:
  1. Put api_key.txt (Gemini key) next to this script
  2. python AutoImageBuilder_v17.py
  3. Select File  OR  Process Folder
  Output saved in Completed_Units/ next to your .docx
"""

# ─── standard library ──────────────────────────────────────────────────────
import os, re, json, textwrap, io, warnings, threading, time

# ─── science / rendering ───────────────────────────────────────────────────
import numpy as np
import matplotlib
matplotlib.use('Agg')                    # non-interactive — safe on any thread
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import (FancyArrowPatch, Arc, Ellipse,
                                 Wedge, FancyBboxPatch)
from mpl_toolkits.mplot3d import Axes3D
from mpl_toolkits.mplot3d.art3d import Poly3DCollection

# ─── document / API ────────────────────────────────────────────────────────
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─── GUI ───────────────────────────────────────────────────────────────────
import customtkinter as ctk
from tkinter import filedialog, messagebox
from datetime import datetime


# ══════════════════════════════════════════════════════════════════════════════
#  CONSTANTS
# ══════════════════════════════════════════════════════════════════════════════
BG_COLOR    = "#F0F8FF"   # gentle light blue  — easy on the eyes
DARK_BLUE   = "#00008B"   # axes, grid, primary curves
RED_ACCENT  = "#CC0000"   # key points, critical values, endpoints
GREEN_LINE  = "#1A7A3C"   # tangent lines, secondary curves
EQ_COLOR    = "#000000"   # pure black — maximum contrast for equation labels
DPI         = 450         # Ultra-HD
FIG_W       = 6.5         # figure width  (inches)
FIG_H       = 5.5         # figure height (inches)
INSERT_W    = 4.5         # width inserted into Word (inches)

MODELS = [
    "gemini-2.0-flash",
    "gemini-2.0-flash-lite",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]


# ══════════════════════════════════════════════════════════════════════════════
#  API KEY  ─ looks next to script, then cwd
# ══════════════════════════════════════════════════════════════════════════════
def load_api_key():
    here = os.path.dirname(os.path.abspath(__file__))
    for p in [os.path.join(here, "api_key.txt"), "api_key.txt"]:
        try:
            k = open(p).read().strip()
            if k: return k
        except: pass
    return None

KEY    = load_api_key()
client = genai.Client(api_key=KEY) if KEY else None


# ══════════════════════════════════════════════════════════════════════════════
#  OMML → PLAIN TEXT
#  p.text silently drops <m:oMath> nodes so AI gets blank equations.
#  This extractor reads every math element and converts it to ASCII math.
# ══════════════════════════════════════════════════════════════════════════════
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _omml(elem):
    tag = elem.tag.split("}")[1] if "}" in elem.tag else elem.tag
    def c(e):
        if e is None: return ""
        return "".join(_omml(ch) for ch in e)
    if tag in ("oMath", "oMathPara"): return c(elem)
    if tag == "r":
        return "".join(t.text or "" for t in elem.findall("{"+MATH_NS+"}t"))
    if tag == "sSup":
        return c(elem.find("{"+MATH_NS+"}e")) + "^" + c(elem.find("{"+MATH_NS+"}sup"))
    if tag == "sSub":
        return c(elem.find("{"+MATH_NS+"}e")) + "_" + c(elem.find("{"+MATH_NS+"}sub"))
    if tag == "sSubSup":
        return (c(elem.find("{"+MATH_NS+"}e")) + "_" +
                c(elem.find("{"+MATH_NS+"}sub")) + "^" +
                c(elem.find("{"+MATH_NS+"}sup")))
    if tag == "f":
        return ("(" + c(elem.find("{"+MATH_NS+"}num")) +
                ")/(" + c(elem.find("{"+MATH_NS+"}den")) + ")")
    if tag == "rad":
        inner = c(elem.find("{"+MATH_NS+"}e"))
        dn    = elem.find("{"+MATH_NS+"}deg")
        dv    = c(dn)
        return ("("+inner+")^(1/"+dv+")") if (dn is not None and dv.strip()) \
               else ("sqrt("+inner+")")
    if tag == "d":
        return "("+",".join(c(e) for e in elem.findall("{"+MATH_NS+"}e"))+")"
    if tag == "func":
        return c(elem.find("{"+MATH_NS+"}fName"))+"("+c(elem.find("{"+MATH_NS+"}e"))+")"
    if tag == "nary":
        ch = elem.find(".//{"+MATH_NS+"}chr")
        sym = ch.get("{"+MATH_NS+"}val","∫") if ch is not None else "∫"
        return (sym+"_"+c(elem.find("{"+MATH_NS+"}sub"))
                +"^"+c(elem.find("{"+MATH_NS+"}sup"))
                +" "+c(elem.find("{"+MATH_NS+"}e")))
    if tag == "limLow":
        return c(elem.find("{"+MATH_NS+"}e"))+"_"+c(elem.find("{"+MATH_NS+"}lim"))
    return "".join(_omml(ch) for ch in elem)

def para_text(p):
    """Full paragraph text including embedded Word math equations."""
    parts = []
    for child in p._p:
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if tag == "r":
            parts.append("".join(t.text or "" for t in child.findall("{"+W_NS+"}t")))
        elif tag in ("oMath","oMathPara"):
            parts.append(_omml(child))
        elif tag in ("hyperlink","ins","del"):
            for r in child.findall(".//{"+W_NS+"}r"):
                parts.append("".join(t.text or "" for t in r.findall("{"+W_NS+"}t")))
    return "".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
#  JSON SANITISER  ─ fixes bad backslash escapes Gemini sometimes produces
# ══════════════════════════════════════════════════════════════════════════════
def sanitize_json(s):
    out = []; i = 0
    while i < len(s):
        if s[i] == "\\" and i+1 < len(s):
            if s[i+1] in ('"',"\\","/","b","f","n","r","t","u"):
                out.append(s[i]); out.append(s[i+1]); i += 2
            else:
                out.append("\\\\"); i += 1
        else:
            out.append(s[i]); i += 1
    return "".join(out)


# ══════════════════════════════════════════════════════════════════════════════
#  CODE SANITISER  ─ fixes common AI matplotlib mistakes before exec()
# ══════════════════════════════════════════════════════════════════════════════
def sanitize_code(code):
    # Fix \\n → real newlines if AI serialised code as single line
    if code and "\n" not in code and "\\n" in code:
        code = code.replace("\\n", "\n")

    # Redirect plt.xxx / ax.property → ax.set_xxx
    _subs = [
        (r'\bplt\.title\s*\(',        'ax.set_title('),
        (r'\bplt\.xlabel\s*\(',       'ax.set_xlabel('),
        (r'\bplt\.ylabel\s*\(',       'ax.set_ylabel('),
        (r'\bplt\.xlim\s*\(',         'ax.set_xlim('),
        (r'\bplt\.ylim\s*\(',         'ax.set_ylim('),
        (r'\bplt\.legend\s*\(',       'ax.legend('),
        (r'\bplt\.grid\s*\(',         'ax.grid('),
        (r'\bplt\.xticks\s*\(',       'ax.set_xticks('),
        (r'\bplt\.yticks\s*\(',       'ax.set_yticks('),
        (r'\bplt\.axhline\s*\(',      'ax.axhline('),
        (r'\bplt\.axvline\s*\(',      'ax.axvline('),
        (r'\bplt\.scatter\s*\(',      'ax.scatter('),
        (r'\bplt\.plot\s*\(',         'ax.plot('),
        (r'\bplt\.fill_between\s*\(', 'ax.fill_between('),
        (r'\bplt\.annotate\s*\(',     'ax.annotate('),
        (r'\bplt\.text\s*\(',         'ax.text('),
        (r'\bax\.title\s*\(',         'ax.set_title('),
        (r'\bax\.xlabel\s*\(',        'ax.set_xlabel('),
        (r'\bax\.ylabel\s*\(',        'ax.set_ylabel('),
    ]
    for pat, rep in _subs:
        code = re.sub(pat, rep, code)

    # Fix invalid escape sequences inside Python string literals
    _valid = set('"\\\'abfnrtvx01234567uUN')
    def _fix(m):
        s = m.group(0); out = []; i = 0
        while i < len(s):
            if s[i]=='\\' and i+1<len(s) and s[i+1] not in _valid:
                out.append('\\\\'); i += 1
            else:
                out.append(s[i]); i += 1
        return "".join(out)
    code = re.sub(r'"(?:[^"\\]|\\.)*"', _fix, code)
    code = re.sub(r"'(?:[^'\\]|\\.)*'", _fix, code)

    # Guard sqrt / log against negative domains
    code = re.sub(r'\bnp\.sqrt\s*\(([^()]+)\)',
                  r'np.sqrt(np.maximum(0,\1))', code)
    code = re.sub(r'\bnp\.log\s*\(([^()]+)\)',
                  r'np.log(np.maximum(1e-12,\1))', code)
    code = re.sub(r'\bnp\.log10\s*\(([^()]+)\)',
                  r'np.log10(np.maximum(1e-12,\1))', code)

    # Strip GUI calls that would crash the process
    for bad in ["plt.show()", "canvas.draw", "draw_idle", "FigureCanvas"]:
        code = code.replace(bad, f"# removed: {bad}")
    return code


# ══════════════════════════════════════════════════════════════════════════════
#  BLANK IMAGE DETECTOR  ─ True if >97 % pixels match background
# ══════════════════════════════════════════════════════════════════════════════
def image_is_blank(path):
    try:
        from PIL import Image
        img  = Image.open(path).convert('RGB')
        arr  = np.array(img, dtype=float)
        diff = np.abs(arr - [240, 248, 255]).mean(axis=2)
        return (diff < 10).mean() > 0.97
    except:
        return False   # PIL unavailable — assume OK


# ══════════════════════════════════════════════════════════════════════════════
#  PROMPTS
# ══════════════════════════════════════════════════════════════════════════════
SYSTEM_INSTRUCTION = r"""
You are a Harvard-level calculus professor AND a precise matplotlib illustrator
producing a professional e-book study guide.

━━━ OUTPUT FORMAT ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Return ONLY a raw JSON array — no markdown, no explanation, nothing else.
Every element must have EXACTLY these 4 fields:

[
  {
    "marker":      "img_001",
    "para_index":  <integer — the [N] index this image follows>,
    "code":        "<complete matplotlib python — \\n between lines>",
    "description": "<4-6 lines — see format below>"
  },
  ...
]

━━━ WHAT TO ILLUSTRATE ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Generate 10–15 images covering:
  • Every theorem statement  (MVT, Rolle's, FTC, IVT, chain rule, etc.)
  • Every worked example     — plot the EXACT function from the text
  • Every practice problem   — geometry diagram OR function graph
  • Key concept explanations (secant/tangent, concavity, critical points)
  • Step-by-step process diagrams where helpful

━━━ CODE RULES (CRITICAL — violations cause blank images) ━━━━━━━━━━━━━━━━━━
Available names: plt  np  ax  fig  matplotlib  mpatches
                 FancyArrowPatch  Arc  Ellipse  Wedge  FancyBboxPatch

NEVER CALL: plt.figure()  plt.show()  fig.add_axes()  fig.add_subplot()
            canvas  FigureCanvas  draw_idle  blit

ALWAYS use ax.plot(), ax.scatter(), ax.text(), ax.annotate(), etc.
Do NOT use plt.legend() — label curves with ax.text() directly on the line.

━━━ VISUAL QUALITY STANDARDS ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Background:    already set to #F0F8FF — do not change
Curves:        linewidth=4.0, color='#00008B'
Key points:    ax.scatter([x],[y], color='#CC0000', s=140, zorder=7)
Equation box on curve (REQUIRED on every image):
  ax.text(x, y, 'f(x) = <exact expression>',
          fontsize=14, fontweight='bold', color='#000000',
          bbox=dict(boxstyle='round,pad=0.5', facecolor='white',
                    edgecolor='#000000', linewidth=2.0, alpha=0.95))
Axis labels:
  ax.set_xlabel('x',  fontsize=14, fontweight='bold', color='#00008B')
  ax.set_ylabel('y',  fontsize=14, fontweight='bold', color='#00008B')
Title:
  ax.set_title('Clear descriptive title',
               fontsize=15, fontweight='bold', color='#00008B', pad=12)
Grid:
  ax.grid(True, alpha=0.25, color='#00008B', linewidth=0.8)
Spines + ticks:
  for s in ax.spines.values(): s.set_color('#00008B'); s.set_linewidth(1.5)
  ax.tick_params(colors='#00008B', labelsize=11)
Margins:  ax.margins(0.15)

━━━ CALCULUS GRAPH TEMPLATE (for MVT, derivatives, function problems) ━━━━━━
# 1. Define the EXACT function — never a placeholder
def f(x): return <exact expression from text, e.g. x**3 - 3*x + 2>
a, b = <exact interval endpoints from text>

# 2. Curve
x_arr = np.linspace(a - 0.4, b + 0.4, 600)
ax.plot(x_arr, f(x_arr), color='#00008B', linewidth=4.0)

# 3. Secant line (for MVT problems)
slope_sec = (f(b) - f(a)) / (b - a)
x_line = np.linspace(a - 0.3, b + 0.3, 100)
ax.plot(x_line, f(a) + slope_sec*(x_line - a),
        color='#CC0000', linewidth=2.5, linestyle='--')

# 4. Tangent at c (for MVT — c is the guaranteed point)
# c_val = <value from problem, e.g. 0.5>
# ax.plot(x_line, f(c_val) + slope_sec*(x_line - c_val),
#         color='#1A7A3C', linewidth=2.5, linestyle='-.')

# 5. Mark and annotate key points
ax.scatter([a, b], [f(a), f(b)], color='#CC0000', s=140, zorder=7)
ax.annotate(f'({a}, f({a}))',
    xy=(a, f(a)), xytext=(a - 0.35, f(a) + 0.4),
    fontsize=11, fontweight='bold', color='#CC0000',
    arrowprops=dict(arrowstyle='->', color='#CC0000', lw=1.5))

# 6. Equation box on curve
mid_x = (a + b) / 2
ax.text(mid_x, f(mid_x) + 0.6, 'f(x) = <expression>',
        fontsize=14, fontweight='bold', color='#000000',
        bbox=dict(boxstyle='round,pad=0.5', facecolor='white',
                  edgecolor='#000000', linewidth=2.0, alpha=0.95))

# 7. Styling
ax.set_title('Title', fontsize=15, fontweight='bold', color='#00008B', pad=12)
ax.set_xlabel('x', fontsize=14, fontweight='bold', color='#00008B')
ax.set_ylabel('y', fontsize=14, fontweight='bold', color='#00008B')
ax.grid(True, alpha=0.25, color='#00008B', linewidth=0.8)
for s in ax.spines.values(): s.set_color('#00008B'); s.set_linewidth(1.5)
ax.tick_params(colors='#00008B', labelsize=11)
ax.margins(0.15)

━━━ GEOMETRY TEMPLATE (optimization, physical objects with dimensions) ━━━━━
ax.set_aspect('equal')
ax.set_xlim(0, total_width); ax.set_ylim(0, total_height)
for s in ax.spines.values(): s.set_visible(False)
ax.set_xticks([]); ax.set_yticks([])

# Shape with fill
shape = plt.Rectangle((x, y), w, h, fc='#AED6F1', ec='#00008B', lw=2.5)
ax.add_patch(shape)

# Double-headed dimension arrows
ax.annotate('', xy=(x2, y_dim), xytext=(x1, y_dim),
    arrowprops=dict(arrowstyle='<->', color='#00008B', lw=2.0))

# All labels in white boxes
ax.text(cx, cy, 'x', fontsize=14, fontweight='bold', ha='center',
    color='#000000',
    bbox=dict(boxstyle='round,pad=0.4', facecolor='white',
              edgecolor='#000000', linewidth=1.8, alpha=0.95))

# Constraint equation at bottom
ax.text(0.5, 0.04, 'Constraint: 2x + y = 1000 m',
    transform=ax.transAxes, fontsize=13, fontweight='bold',
    ha='center', color='#000000',
    bbox=dict(boxstyle='round,pad=0.4', facecolor='#FFFDE7',
              edgecolor='#000000', linewidth=1.8))

ax.set_title('Title', fontsize=15, fontweight='bold', color='#00008B', pad=12)
ax.margins(0.15)

━━━ DESCRIPTION FORMAT ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Line 1: Section / problem reference (e.g. "Section 5.1.3, Example 2")
Line 2: Concept being illustrated (e.g. "Mean Value Theorem applied to f(x)=x²+2x−1 on [0,1]")
Line 3: Mathematical explanation of what is shown in the image
Line 4: What each labeled point, line, or region represents
Line 5: What a student should observe or conclude from this image
Line 6: (optional) Connection to the theorem or solution step
"""

USER_PROMPT = """Create ULTRA-HD professional calculus visualizations for this document.

REQUIREMENTS:
- 10-15 images total: theorems, worked examples, practice problems
- ALL 4 fields per image: marker, para_index, code, description  
- Plot the EXACT function from the text — never invent or simplify
- Equation labels in pure black (#000000) with white rounded boxes directly on curves
- 4-6 line descriptions per image
- Use ax.text() on curves, NOT plt.legend()

DOCUMENT:
{content}"""


# ══════════════════════════════════════════════════════════════════════════════
#  WORKER  ─ runs entirely on a background thread
# ══════════════════════════════════════════════════════════════════════════════
class Worker:
    def __init__(self, log_cb, progress_cb, done_cb):
        self._log  = log_cb       # fn(str)   — thread-safe via after(0,...)
        self._prog = progress_cb  # fn(float) — 0.0–1.0
        self._done = done_cb      # fn(bool, str)

    def log(self, msg):      self._log(msg)
    def prog(self, v):       self._prog(v)

    # ── public entry points ───────────────────────────────────
    def run_file(self, path):
        try:
            ok, fail, saved = self._process(path)
            self._done(True,
                f"Complete!\n\nSaved to:\n{saved}\n\n"
                f"Images inserted: {ok}   Failed/skipped: {fail}")
        except Exception as e:
            import traceback
            self.log(f"\n❌ {e}\n{traceback.format_exc()[:900]}")
            self._done(False, str(e))

    def run_folder(self, folder):
        files = sorted([
            os.path.join(folder, f) for f in os.listdir(folder)
            if f.endswith(".docx") and not f.startswith("~")
        ])
        if not files:
            self._done(False, "No .docx files found in that folder.")
            return
        self.log(f"Folder queue: {len(files)} file(s)")
        total_ok = 0
        for i, path in enumerate(files, 1):
            self.prog((i - 1) / len(files))
            self.log(f"\n{'='*55}")
            self.log(f"[{i}/{len(files)}] {os.path.basename(path)}")
            self.log(f"{'='*55}")
            try:
                ok, fail, saved = self._process(path)
                total_ok += ok
                self.log(f"✅  {ok} images  →  {os.path.basename(saved)}")
            except Exception as e:
                self.log(f"❌  {e}")
        self.prog(1.0)
        self._done(True,
            f"Folder complete!\n{len(files)} files processed.\n"
            f"Total images inserted: {total_ok}")

    # ── core processing ───────────────────────────────────────
    def _process(self, doc_path):
        log = self.log

        # ── 1. Load and extract text ──────────────────────────
        log("Loading document…")
        doc   = Document(doc_path)
        lines = []
        for i, p in enumerate(doc.paragraphs):
            t = para_text(p).strip()
            if t:
                lines.append(f"[{i}] {t}")
        content = "\n".join(lines)
        log(f"  {len(doc.paragraphs)} paragraphs, {len(lines)} non-empty")

        out_dir = os.path.join(os.path.dirname(doc_path), "Completed_Units")
        os.makedirs(out_dir, exist_ok=True)

        # ── 2. ONE AI call for the whole document ─────────────
        log("\n── AI analysis (one call, up to 3 min) ──")
        t0  = time.time()
        raw = None
        for model in MODELS:
            try:
                log(f"  Trying {model}…")
                r = client.models.generate_content(
                    model=model,
                    contents=USER_PROMPT.replace("{content}", content),
                    config={
                        "system_instruction": SYSTEM_INSTRUCTION,
                        "temperature": 0.2,
                    })
                raw = r.text
                log(f"  ✓ {model} replied in {time.time()-t0:.1f}s  "
                    f"({len(raw)} chars)")
                break
            except Exception as e:
                log(f"  ✗ {model}: {e}")

        if not raw:
            raise RuntimeError(
                "All AI models failed. "
                "Check api_key.txt and your internet connection.")
        self.prog(0.30)

        # ── 3. Parse JSON ─────────────────────────────────────
        log("  Parsing JSON…")
        clean = re.sub(r'^```[a-zA-Z]*\s*', '', raw.strip(), flags=re.MULTILINE)
        clean = re.sub(r'```\s*$', '', clean.strip())
        arr_m = re.search(r'\[.*\]', clean, re.DOTALL)
        if not arr_m:
            raise RuntimeError(
                f"AI returned no JSON array.\nFirst 400 chars:\n{raw[:400]}")

        json_str = sanitize_json(arr_m.group())
        tasks = []
        try:
            tasks = json.loads(json_str)
        except json.JSONDecodeError:
            log("  Standard parse failed — extracting objects one by one…")
            for obj in re.findall(r'\{[^{}]{30,}\}', json_str, re.DOTALL):
                try: tasks.append(json.loads(sanitize_json(obj)))
                except: pass

        # Validate tasks
        valid = []
        for t in tasks:
            if not isinstance(t, dict): continue
            if not all(k in t for k in ('marker','code','description','para_index')):
                continue
            try:
                t['para_index'] = int(t['para_index'])
                valid.append(t)
            except: pass

        log(f"  {len(valid)} valid image tasks")
        if not valid:
            raise RuntimeError("No valid tasks returned by AI.")
        self.prog(0.38)

        # ── 4. Render each image ──────────────────────────────
        log(f"\n── Rendering {len(valid)} Ultra-HD images (450 DPI) ──")
        image_data = []
        n_fail     = 0

        for idx, task in enumerate(valid):
            pid  = task['para_index']
            desc = task.get('description', '').strip()
            code = task.get('code', '')

            log(f"  [{idx+1}/{len(valid)}] para={pid}")

            # Block GUI-crash calls
            if any(bad.lower() in code.lower()
                   for bad in ('canvas','FigureCanvas','draw_idle','blit')):
                log("    ✗ forbidden call"); n_fail += 1; continue

            # Sanitise
            code = sanitize_code(code)

            # Syntax check
            try:
                compile(code, "<check>", "exec")
            except SyntaxError as se:
                log(f"    ✗ SyntaxError: {se}"); n_fail += 1; continue

            # Unique temp filename
            tmp = os.path.join(
                out_dir,
                f"_tmp_{pid}_{int(time.time()*1_000_000)}.png")

            rendered = False
            for attempt in range(1, 3):           # up to 2 attempts
                try:
                    fig = plt.figure(figsize=(FIG_W, FIG_H), dpi=DPI)
                    ax  = fig.add_axes([0.14, 0.14, 0.82, 0.82])
                    ax.set_facecolor(BG_COLOR)
                    fig.patch.set_facecolor(BG_COLOR)

                    scope = {
                        "plt": plt, "np": np, "ax": ax, "fig": fig,
                        "matplotlib": matplotlib, "mpatches": mpatches,
                        "FancyArrowPatch": FancyArrowPatch, "Arc": Arc,
                        "Ellipse": Ellipse, "Wedge": Wedge,
                        "FancyBboxPatch": FancyBboxPatch,
                        "Axes3D": Axes3D,
                        "Poly3DCollection": Poly3DCollection,
                        "__builtins__": __builtins__,
                    }
                    with warnings.catch_warnings():
                        warnings.simplefilter("ignore")
                        exec(compile(code, "<ai>", "exec"), scope)

                    plt.savefig(tmp, facecolor=BG_COLOR,
                                bbox_inches='tight', pad_inches=0.25,
                                dpi=DPI, format='png')
                    plt.close('all')

                    # Verify file was created and is non-blank
                    if not os.path.exists(tmp) or os.path.getsize(tmp) < 1000:
                        raise Exception("Image file too small or missing")
                    if image_is_blank(tmp):
                        raise Exception("Image rendered blank — nothing drawn")

                    kb = os.path.getsize(tmp) // 1024
                    image_data.append(
                        {'para_index': pid, 'img_name': tmp, 'description': desc})
                    log(f"    ✓ {kb} KB")
                    rendered = True
                    break

                except Exception as e:
                    plt.close('all')
                    if os.path.exists(tmp):
                        try: os.remove(tmp)
                        except: pass
                    log(f"    ✗ attempt {attempt}: {str(e)[:90]}")
                    # On retry: apply sanitiser again more aggressively
                    code = sanitize_code(code)

            if not rendered:
                n_fail += 1

            self.prog(0.38 + 0.47 * (idx + 1) / len(valid))

        log(f"\n── {len(image_data)} rendered,  {n_fail} failed ──")
        if not image_data:
            raise RuntimeError(
                f"No images rendered successfully. {n_fail} failed.")

        # ── 5. Insert images into Word document ──────────────
        log(f"── Inserting {len(image_data)} images ──")
        n_ok    = 0
        n_ins_f = 0
        n_paras = len(doc.paragraphs)

        # Insert in reverse order so earlier indices stay valid
        for img in sorted(image_data,
                          key=lambda x: x['para_index'], reverse=True):
            pid  = img['para_index']
            path = img['img_name']
            desc = img['description']

            if pid < 0 or pid >= n_paras:
                log(f"  skip para={pid} (out of range)"); continue
            if not os.path.exists(path):
                log(f"  skip para={pid} (file missing)"); continue

            try:
                target = doc.paragraphs[pid]

                # Image paragraph — centered
                img_para = doc.add_paragraph()
                target._p.addnext(img_para._p)
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(path, width=Inches(INSERT_W))

                # Description paragraph — 4-6 lines, dark blue italic
                if len(desc) >= 50:
                    wrapped = textwrap.wrap(desc, width=80)
                    desc_para = doc.add_paragraph()
                    img_para._p.addnext(desc_para._p)
                    run = desc_para.add_run("\n".join(wrapped))
                    run.font.size      = Pt(10)
                    run.font.italic    = True
                    run.font.color.rgb = RGBColor(0, 0, 139)
                    run.font.name      = "Calibri"
                    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    desc_para.paragraph_format.space_after = Pt(12)

                n_ok += 1
                log(f"  ✓ para={pid}")

            except Exception as e:
                log(f"  ✗ para={pid}: {e}")
                n_ins_f += 1

        # ── 6. Save ───────────────────────────────────────────
        save_name = os.path.basename(doc_path).replace(".docx", "_Professional.docx")
        save_path = os.path.join(out_dir, save_name)
        log(f"\nSaving  →  {save_path}")
        doc.save(save_path)
        log("Saved ✓")

        # ── 7. Clean up temp PNGs ─────────────────────────────
        time.sleep(1)
        for img in image_data:
            try:
                if os.path.exists(img['img_name']):
                    os.remove(img['img_name'])
            except: pass
        log("Temp files cleaned ✓")

        self.prog(1.0)
        return n_ok, n_fail + n_ins_f, save_path


# ══════════════════════════════════════════════════════════════════════════════
#  GUI  ─ main thread only — never does any processing itself
# ══════════════════════════════════════════════════════════════════════════════
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("AutoImageBuilder  v17  —  Professional")
        self.geometry("800x740")
        self._thread = None

        # ── Header ────────────────────────────────────────────
        hdr = ctk.CTkFrame(self, fg_color="#0D1B2A", corner_radius=12)
        hdr.pack(pady=14, padx=20, fill="x")
        ctk.CTkLabel(hdr, text="AutoImageBuilder  v17",
                     font=("Georgia", 28, "bold"),
                     text_color="#F0C040").pack(pady=(14, 2))
        ctk.CTkLabel(hdr,
                     text="450 DPI · Deep Parsing · Accurate Diagrams · Background Thread",
                     font=("Arial", 11),
                     text_color="#7FB3D3").pack(pady=(0, 14))

        # ── Feature card ──────────────────────────────────────
        card = ctk.CTkFrame(self, fg_color="#EEF4FF", corner_radius=10)
        card.pack(pady=6, padx=20, fill="x")
        ctk.CTkLabel(card, text="What this tool does",
                     font=("Arial", 12, "bold"),
                     text_color="#0D1B2A").pack(pady=(10, 4))
        for ln in [
            "🧮  Reads Word math equations (OMML) — exact functions, never guesses",
            "📈  Accurate calculus graphs: exact curve, secant + tangent, labeled points",
            "📐  Accurate geometry: labeled dimensions, constraint equations, arrows",
            "👁️  450 DPI Ultra-HD — razor-sharp at 4.5\" wide in Word",
            "📝  4-6 line description under every image: concept, math, insight",
            "⚡  One AI call per document — 1 to 3 minutes total",
            "🔒  Background thread — window stays fully responsive",
            "🔍  Blank image detection — auto-retry if render produces nothing",
            "📁  Folder mode processes every .docx file in one click",
        ]:
            ctk.CTkLabel(card, text=ln, font=("Arial", 10),
                         text_color="#1a2a3a").pack(
                pady=1, anchor="w", padx=16)
        ctk.CTkLabel(card, text="",
                     font=("Arial", 4)).pack()  # spacer

        # ── API key status ────────────────────────────────────
        kf = ctk.CTkFrame(self,
                          fg_color="#FFF3E0" if not KEY else "#E8F5E9",
                          corner_radius=8)
        kf.pack(pady=(8, 2), padx=20, fill="x")
        ctk.CTkLabel(kf,
                     text=("⚠️  api_key.txt not found — place it next to this script"
                           if not KEY
                           else f"✅  Gemini API key loaded  ({len(KEY)} chars)"),
                     font=("Arial", 10, "bold"),
                     text_color="#BF360C" if not KEY else "#1B5E20").pack(pady=7)

        # ── Buttons ───────────────────────────────────────────
        self.btn_file = ctk.CTkButton(
            self, text="📄   Select File & Process",
            command=self.pick_file, height=54, width=450,
            font=("Arial", 15, "bold"),
            fg_color="#1565C0", hover_color="#0D47A1")
        self.btn_file.pack(pady=(10, 4))

        self.btn_folder = ctk.CTkButton(
            self, text="📁   Process Entire Folder",
            command=self.pick_folder, height=54, width=450,
            font=("Arial", 15, "bold"),
            fg_color="#1B5E20", hover_color="#0A3010")
        self.btn_folder.pack(pady=4)

        # ── Progress bar ──────────────────────────────────────
        self.prog_bar = ctk.CTkProgressBar(self, width=450, height=14)
        self.prog_bar.set(0)
        self.prog_bar.pack(pady=8)

        # ── Status label ──────────────────────────────────────
        sf = ctk.CTkFrame(self, fg_color="#E3F2FD", corner_radius=8)
        sf.pack(pady=4, padx=20, fill="x")
        self.status_lbl = ctk.CTkLabel(
            sf, text="● Ready",
            text_color="#1565C0", font=("Arial", 12, "bold"))
        self.status_lbl.pack(pady=7)

        # ── Console ───────────────────────────────────────────
        ctk.CTkLabel(self, text="Console (live progress):",
                     font=("Arial", 10, "bold")).pack(
            pady=(4, 0), padx=20, anchor="w")
        self.con = ctk.CTkTextbox(
            self, height=210, font=("Courier", 9),
            fg_color="#0D1B2A", text_color="#B3D4F5")
        self.con.pack(pady=(2, 10), padx=20, fill="both", expand=True)

    # ── Thread-safe UI helpers ─────────────────────────────────
    def _log(self, msg):
        def _do():
            ts = datetime.now().strftime("%H:%M:%S")
            self.con.insert("end", f"[{ts}] {msg}\n")
            self.con.see("end")
        self.after(0, _do)

    def _set_prog(self, v):
        self.after(0, lambda: self.prog_bar.set(v))

    def _set_status(self, msg, color):
        self.after(0, lambda: self.status_lbl.configure(
            text=msg, text_color=color))

    def _set_btns(self, enabled):
        s = "normal" if enabled else "disabled"
        self.after(0, lambda: self.btn_file.configure(state=s))
        self.after(0, lambda: self.btn_folder.configure(state=s))

    def _on_done(self, ok, msg):
        self._set_btns(True)
        self._set_status("● Ready", "#1565C0")
        if ok:
            self.after(0, lambda: messagebox.showinfo("Complete", msg))
        else:
            self.after(0, lambda: messagebox.showerror("Error", msg))

    # ── Busy guard ────────────────────────────────────────────
    def _busy(self):
        if self._thread and self._thread.is_alive():
            messagebox.showwarning("Busy", "Still processing — please wait.")
            return True
        return False

    # ── Launch helper ─────────────────────────────────────────
    def _launch(self, method, arg):
        self.con.delete("1.0", "end")
        self._set_status("● Working…", "#E65100")
        self._set_prog(0)
        self._set_btns(False)
        w = Worker(
            log_cb=self._log,
            progress_cb=self._set_prog,
            done_cb=self._on_done)
        self._thread = threading.Thread(
            target=method, args=(w, arg), daemon=True)
        self._thread.start()

    # ── Button handlers ───────────────────────────────────────
    def pick_file(self):
        if self._busy(): return
        if not KEY:
            messagebox.showerror("No API Key",
                "Create a file called api_key.txt containing your Gemini API key\n"
                "and place it in the same folder as this script."); return
        path = filedialog.askopenfilename(
            title="Select Word document",
            filetypes=[("Word Documents", "*.docx")])
        if path:
            self._launch(lambda w, p: w.run_file(p), path)

    def pick_folder(self):
        if self._busy(): return
        if not KEY:
            messagebox.showerror("No API Key",
                "Create a file called api_key.txt containing your Gemini API key\n"
                "and place it in the same folder as this script."); return
        folder = filedialog.askdirectory(title="Select folder of .docx files")
        if folder:
            self._launch(lambda w, f: w.run_folder(f), folder)


# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    App().mainloop()
