import os
import re
import json
import textwrap
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from google import genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import customtkinter as ctk
from tkinter import filedialog, messagebox
from datetime import datetime

# --- MASTER CONFIGURATION ---
BG_COLOR = "#F0F8FF"    # Very Light Blue
DARK_BLUE = "#00008B"   # Sharp Dark Blue
RED_ACCENT = "#FF0000"  # For Holes/Asymptotes
EQUATION_COLOR = "#000000"  # PURE BLACK for maximum contrast on equations

def load_api_key():
    try:
        with open("api_key.txt", "r") as f:
            return f.read().strip()
    except FileNotFoundError:
        return None

client = genai.Client(api_key=load_api_key())

# --- SYSTEM INSTRUCTION WITH STRICT JSON FORMAT ---
SYSTEM_INSTRUCTION = r"""
Role: Harvard Math Professor & Technical Illustrator for Professional E-books.

🎯 MANDATORY JSON FORMAT - ALL FIELDS REQUIRED:

{
  "marker": "unique_id_string",
  "code": "complete matplotlib code",
  "description": "4-6 line detailed description",
  "para_index": numeric_paragraph_index
}

ALL FOUR FIELDS ARE MANDATORY. Missing any field = INVALID.

═══════════════════════════════════════════════════════════════════════════════
REQUIREMENTS:
═══════════════════════════════════════════════════════════════════════════════

1. MINIMUM 15-20 images per document
2. EXACT functions from document (never modify)
3. Descriptions MUST be 4-6 lines with detailed mathematical explanation
4. Equations placed DIRECTLY on curves using plt.text()
5. ULTRA-HD quality: DPI 600, antialiasing, sharp crisp text
6. Equations in PURE BLACK (#000000) for maximum contrast
7. JSON SAFETY: Use ONLY standard ASCII - NO Unicode symbols (∞, ≤, π)

JSON SAFETY CRITICAL RULES:
- All JSON strings MUST use standard ASCII characters only
- NO Unicode math symbols - write them as words or use standard characters
- NO \x escape sequences in any strings
- Properly escape special characters: \\n for newlines, \\" for quotes
- Test that JSON is valid before returning

═══════════════════════════════════════════════════════════════════════════════
EQUATION PLACEMENT ON CURVES - BOLD & HIGH CONTRAST:
═══════════════════════════════════════════════════════════════════════════════

✅ CORRECT - BOLD BLACK equations with high contrast:
```python
plt.text(x_pos, y_pos, 'f(x)=x²', 
         fontsize=15, weight='bold', color='#000000',  # PURE BLACK
         bbox=dict(boxstyle='round,pad=0.5', facecolor='white', 
                   edgecolor='#000000', linewidth=2.0, alpha=0.95))
```

CRITICAL: 
- Equation color MUST be #000000 (pure black) for maximum contrast
- Font weight MUST be 'bold'
- Fontsize 14-16 for equations
- White background with black border for readability

❌ WRONG - Don't use dim colors or legends:
plt.text(..., color='#00008B')  # NO! Too dim
plt.legend()  # NO!

═══════════════════════════════════════════════════════════════════════════════
ULTRA-HD PROFESSIONAL QUALITY:
═══════════════════════════════════════════════════════════════════════════════

- DPI: 600 (ultra-sharp, print quality)
- Antialiasing: ON (smooth text rendering)
- Background: #F0F8FF (gentle on eyes)
- Axes/Grid: #00008B (high contrast)
- Equations: #000000 (PURE BLACK, maximum contrast)
- Axes labels: fontsize=14, weight='bold'
- Equations: fontsize=14-16, weight='bold', color='#000000'
- Grid: alpha=0.25, subtle
- Lines: linewidth=4.0, ultra-clear

═══════════════════════════════════════════════════════════════════════════════
DETAILED DESCRIPTIONS (4-6 LINES MINIMUM):
═══════════════════════════════════════════════════════════════════════════════

Each description MUST include:
1. What concept/theorem is being illustrated
2. Mathematical explanation of what the graph shows
3. Key features (intercepts, asymptotes, behavior, etc.)
4. Practical interpretation or significance
5. Connection to the problem or concept in the document

EXAMPLE (4-6 lines):
"Problem 1.2: Quadratic Function and Its Zeros
This graph illustrates the parabola f(x)=x²-3x+2 and demonstrates how factoring reveals the x-intercepts. By factoring as (x-1)(x-2)=0, we identify zeros at x=1 and x=2, shown as red dots on the graph. The parabola opens upward (positive leading coefficient) with vertex between the zeros. This visualization connects algebraic factoring to geometric interpretation, showing that zeros of a function correspond to points where the graph crosses the x-axis."

═══════════════════════════════════════════════════════════════════════════════
EXACT FUNCTIONS - CRITICAL:
═══════════════════════════════════════════════════════════════════════════════

Document: "f(x) = x² + 3x - 4"
Code: f = x**2 + 3*x - 4  ✓ EXACT COPY

Document: "g(x) = (x-2)(x+3)"
Code: g = (x-2)*(x+3)  ✓ EXACT COPY

═══════════════════════════════════════════════════════════════════════════════
COMPLETE EXAMPLE WITH ULTRA-HD & DETAILED DESCRIPTION:
═══════════════════════════════════════════════════════════════════════════════

[
  {
    "marker": "problem_1_2_parabola",
    "code": "import numpy as np\nimport matplotlib.pyplot as plt\n\nx = np.linspace(-2, 5, 800)\ny = x**2 - 3*x + 2\n\nplt.plot(x, y, color='#00008B', linewidth=4.0, antialiased=True)\nplt.text(2.5, 2.0, 'f(x)=x²-3x+2', fontsize=15, weight='bold', color='#000000', bbox=dict(boxstyle='round,pad=0.5', facecolor='white', edgecolor='#000000', linewidth=2.0, alpha=0.95))\n\nplt.scatter([1, 2], [0, 0], s=150, color='#DC143C', zorder=5, edgecolors='#000000', linewidths=1.5)\nplt.text(1, -0.5, '(1,0)', fontsize=12, weight='bold', color='#000000', ha='center')\nplt.text(2, -0.5, '(2,0)', fontsize=12, weight='bold', color='#000000', ha='center')\n\nplt.axhline(0, color='#00008B', linewidth=2.0)\nplt.axvline(0, color='#00008B', linewidth=2.0)\nplt.grid(True, alpha=0.25, color='#00008B', linestyle='-', linewidth=0.8)\nplt.xlabel('x', fontsize=14, weight='bold', color='#00008B')\nplt.ylabel('y', fontsize=14, weight='bold', color='#00008B')\nplt.xlim(-2, 5)\nplt.ylim(-2, 6)",
    "description": "Problem 1.2: Quadratic Function and Its Zeros\nThis graph illustrates the parabola f(x)=x²-3x+2 and demonstrates how factoring reveals the x-intercepts. By factoring as (x-1)(x-2)=0, we identify zeros at x=1 and x=2, shown as red dots on the graph. The parabola opens upward (positive leading coefficient) with vertex at x=1.5. This visualization connects algebraic factoring to geometric interpretation, showing that zeros correspond to x-axis intersection points.",
    "para_index": 12
  }
]

VERIFY ALL 4 FIELDS PRESENT IN EVERY OBJECT.
DESCRIPTIONS MUST BE 4-6 LINES.
EQUATIONS MUST BE PURE BLACK (#000000).
"""

def parse_xml(xml_string):
    """Helper to parse XML string"""
    from lxml import etree
    return etree.fromstring(xml_string)

class STEMOrchestrator:
    def process_file(self, doc_path, image_size=4.5):
        doc = Document(doc_path)
        content = "\n".join([f"[{i}] {p.text}" for i, p in enumerate(doc.paragraphs) if p.text.strip()])
        
        try:
            response = client.models.generate_content(
                model='gemini-3-flash-preview',
                contents=r"""Create ULTRA-HD PROFESSIONAL E-BOOK visualizations.

CRITICAL JSON FORMAT - MUST START WITH [ AND END WITH ]:

[
  {
    "marker": "string",
    "code": "matplotlib code",
    "description": "4-6 lines detailed explanation",
    "para_index": number
  },
  {
    "marker": "string",
    "code": "matplotlib code", 
    "description": "4-6 lines detailed explanation",
    "para_index": number
  }
]

IMPORTANT: Response MUST be a valid JSON array starting with [ and ending with ]

ULTRA-HD REQUIREMENTS:
1. Generate 15-20 images (entire document)
2. Equations ON curves (plt.text), PURE BLACK color (#000000), fontsize 14-16, weight='bold'
3. Exact functions from document
4. ULTRA-SHARP: DPI 600, antialiased=True, linewidth=4.0
5. Descriptions 4-6 lines with detailed mathematical explanation
6. Maximum contrast: equations in #000000 (pure black) on white background

JSON SAFETY RULES (CRITICAL):
- Use ONLY standard ASCII characters in JSON
- NO special math symbols (infinity, less-equal, greater-equal, pi, etc.) - write as words
- NO backslash-x escape sequences in strings
- Use double-backslash-n for newlines in descriptions
- Use simple quotes and apostrophes only
- Test all JSON is valid before returning

EQUATION STYLE (MANDATORY):
plt.text(x, y, 'f(x)=...', fontsize=15, weight='bold', color='#000000',
         bbox=dict(boxstyle='round,pad=0.5', facecolor='white', 
                   edgecolor='#000000', linewidth=2.0, alpha=0.95))

DESCRIPTION MUST INCLUDE:
- Concept being illustrated (1 line)
- Mathematical explanation (2 lines)
- Key features and interpretation (1-2 lines)
- Connection to problem/theorem (1 line)

Document with paragraph indices:
""" + f"""{content}

Return valid JSON array with ALL 4 fields in every object.""",
                config={'system_instruction': SYSTEM_INSTRUCTION}
            )
            
            # Extract JSON from response - handle various formats
            response_text = response.text.strip()
            
            # Try to find JSON array in response
            clean_json = re.search(r'\[.*\]', response_text, re.DOTALL)
            
            # If no complete array found, try to fix common issues
            if not clean_json:
                print("WARNING: No complete JSON array found. Attempting to fix...")
                
                # Check if array is missing opening bracket
                if response_text.strip().startswith('{'):
                    print("  → Missing opening bracket - adding [")
                    response_text = '[' + response_text
                
                # Check if array is missing closing bracket
                if not response_text.strip().endswith(']'):
                    print("  → Missing closing bracket - adding ]")
                    response_text = response_text + ']'
                
                # Try again
                clean_json = re.search(r'\[.*\]', response_text, re.DOTALL)
            
            if not clean_json:
                print("ERROR: No JSON found in AI response")
                print(f"Response preview: {response_text[:500]}")
                # Save for debugging
                with open("debug_response.txt", "w", encoding='utf-8') as f:
                    f.write(response_text)
                print("Saved full response to debug_response.txt")
                return "AI did not return valid JSON"
            
            # Clean the JSON string to handle escape sequences
            json_str = clean_json.group()
            
            # Fix common escape sequence issues
            # Replace problematic escape sequences that aren't valid JSON
            json_str = json_str.replace('\\x', '\\\\x')  # Fix \x escape sequences
            
            # Try to parse with better error reporting
            try:
                tasks = json.loads(json_str)
            except json.JSONDecodeError as e:
                print(f"JSON Parse Error at line {e.lineno}, column {e.colno}")
                print(f"Error message: {e.msg}")
                print(f"\nProblematic section:")
                # Show context around the error
                lines = json_str.split('\n')
                start = max(0, e.lineno - 3)
                end = min(len(lines), e.lineno + 2)
                for i in range(start, end):
                    marker = ">>> " if i == e.lineno - 1 else "    "
                    if i < len(lines):
                        print(f"{marker}Line {i+1}: {lines[i][:100]}")
                
                # Try alternative cleaning methods
                print("\nAttempting aggressive JSON cleaning...")
                
                # Method 2: Remove all problematic escape sequences
                json_str_clean = re.sub(r'\\x[0-9a-fA-F]{2}', '', json_str)
                
                try:
                    tasks = json.loads(json_str_clean)
                    print("✓ Successfully parsed with aggressive cleaning")
                except Exception as e2:
                    # Last resort: save the problematic JSON for debugging
                    with open("debug_json.txt", "w", encoding='utf-8') as f:
                        f.write(json_str)
                    print("ERROR: Could not parse JSON even after cleaning")
                    print("Saved problematic JSON to debug_json.txt")
                    print(f"Final error: {str(e2)}")
                    return f"JSON parsing failed: {e.msg}"
            
            print(f"✓ AI returned {len(tasks)} tasks")
            
            # VALIDATE all required fields
            valid_tasks = []
            for i, task in enumerate(tasks):
                if not all(k in task for k in ['marker', 'code', 'description', 'para_index']):
                    print(f"⚠ Task {i} missing required fields - SKIPPED")
                    print(f"   Has keys: {list(task.keys())}")
                    continue
                
                # Validate para_index is a number
                try:
                    task['para_index'] = int(task['para_index'])
                    valid_tasks.append(task)
                except (ValueError, TypeError):
                    print(f"⚠ Task {i} has invalid para_index: {task.get('para_index')} - SKIPPED")
                    continue
            
            if not valid_tasks:
                print("ERROR: No valid tasks after validation")
                return "AI did not return any valid tasks"
            
            print(f"✓ Validated {len(valid_tasks)} tasks (discarded {len(tasks) - len(valid_tasks)})")
            tasks = valid_tasks
            
            if len(tasks) < 10:
                print(f"⚠ WARNING: Only {len(tasks)} valid images. Target: 15+")
                    
        except json.JSONDecodeError as e:
            print(f"ERROR: JSON parsing failed: {e}")
            return f"AI returned invalid JSON: {str(e)}"
        except Exception as e:
            print(f"ERROR: {e}")
            return f"AI Error: {str(e)}"

        output_dir = os.path.join(os.path.dirname(doc_path), "Completed_Units")
        os.makedirs(output_dir, exist_ok=True)

        successful = 0
        failed = 0
        temp_files_to_delete = []  # CRITICAL: Track temp files for cleanup AFTER document save
        
        for task in sorted(tasks, key=lambda x: x['para_index'], reverse=True):
            img_name = f"temp_{task['para_index']}.png"
            
            # Validate code
            forbidden_terms = ['canvas', 'FigureCanvas', 'draw_idle', 'blit', 'add_subplot']
            if any(term.lower() in task['code'].lower() for term in forbidden_terms):
                print(f"⚠ Skipped para {task['para_index']}: forbidden term")
                failed += 1
                continue
            
            description_text = task.get('description', '').strip()
            has_description = len(description_text) >= 50  # Allow longer descriptions (4-6 lines)
            
            try:
                # ULTRA-HD Professional quality figure - LARGER SIZE
                fig = plt.figure(figsize=(8.0, 6.5), dpi=600)  # Increased from 6.5x5.5 to 8x6.5
                ax = fig.add_axes([0.14, 0.14, 0.82, 0.82])
                ax.set_facecolor(BG_COLOR)
                fig.patch.set_facecolor(BG_COLOR)
                
                # Execute code
                exec_scope = {
                    "plt": plt, "np": np, "matplotlib": matplotlib,
                    "fig": fig, "ax": ax, "__builtins__": __builtins__
                }
                try:
                    exec(task['code'], exec_scope)
                except Exception as exec_err:
                    raise Exception(f"Code execution error: {str(exec_err)}")
                
                # Save ULTRA-HD QUALITY
                try:
                    plt.savefig(img_name, 
                               facecolor=BG_COLOR, 
                               bbox_inches='tight',
                               pad_inches=0.25, 
                               dpi=600,  # ULTRA-HD
                               format='png',
                               metadata={'Software': 'Professional E-Book Builder Ultra-HD'})
                except Exception as save_err:
                    raise Exception(f"Image save error: {str(save_err)}")
                finally:
                    plt.close('all')

                # Insert using STABLE TABLE structure
                if task['para_index'] < len(doc.paragraphs):
                    # Verify image file was created successfully
                    if not os.path.exists(img_name):
                        raise Exception(f"Image file {img_name} was not created")
                    
                    # Verify file is not empty
                    if os.path.getsize(img_name) == 0:
                        raise Exception(f"Image file {img_name} is empty")
                    
                    target_para = doc.paragraphs[task['para_index']]
                    
                    # Create STABLE 2-column table with fixed layout
                    table = doc.add_table(rows=1, cols=2)
                    target_para._p.addnext(table._tbl)
                    
                    # Set table to FIXED layout for stability
                    tbl = table._tbl
                    tblPr = tbl.tblPr
                    if tblPr is None:
                        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                        tbl.insert(0, tblPr)
                    
                    # Add table layout property for stability
                    tblLayout = parse_xml(r'<w:tblLayout w:type="fixed" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    tblPr.append(tblLayout)
                    
                    # Set table width to full page width
                    tblW = parse_xml(r'<w:tblW w:w="9000" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    tblPr.append(tblW)
                    
                    # Remove ALL borders for invisible table
                    table.style = 'Table Grid'
                    for row in table.rows:
                        for cell in row.cells:
                            tc = cell._element
                            tcPr = tc.get_or_add_tcPr()
                            
                            # Remove borders
                            tcBorders = parse_xml(
                                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:val="none"/><w:left w:val="none"/>'
                                r'<w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'
                            )
                            tcPr.append(tcBorders)
                            
                            # Add cell margins for spacing
                            tcMar = parse_xml(
                                r'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                r'<w:top w:w="100" w:type="dxa"/>'
                                r'<w:left w:w="100" w:type="dxa"/>'
                                r'<w:bottom w:w="100" w:type="dxa"/>'
                                r'<w:right w:w="100" w:type="dxa"/>'
                                r'</w:tcMar>'
                            )
                            tcPr.append(tcMar)
                    
                    # Dynamic column widths based on selected image size
                    # Total page width: 7.5 inches
                    # Left column gets remaining space
                    left_width = 7.5 - image_size
                    
                    # Left column: Description (remaining space after image)
                    left_cell = table.rows[0].cells[0]
                    left_cell.width = Inches(left_width)
                    
                    # Set cell width property for stability
                    left_dxa = int(left_width * 1000)  # Convert inches to DXA
                    tcW = parse_xml(f'<w:tcW w:w="{left_dxa}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    left_cell._element.get_or_add_tcPr().append(tcW)
                    
                    # Add vertical alignment
                    vAlign = parse_xml(r'<w:vAlign w:val="top" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    left_cell._element.get_or_add_tcPr().append(vAlign)
                    
                    if has_description:
                        # Adjust wrap width based on column width
                        chars_per_inch = 17  # Approximate characters per inch at 10pt
                        wrap_width = int(left_width * chars_per_inch)
                        
                        # FULL description - NO truncation
                        wrapped_lines = textwrap.wrap(description_text, width=wrap_width)
                        # Don't limit to 6 lines - show FULL description
                        
                        desc_para = left_cell.paragraphs[0]
                        desc_run = desc_para.add_run("\n".join(wrapped_lines))
                        desc_run.font.size = Pt(10)
                        desc_run.font.color.rgb = RGBColor(0, 0, 139)
                        desc_run.font.name = 'Calibri'
                        desc_run.font.bold = False
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        desc_para.paragraph_format.line_spacing = 1.15  # Comfortable reading
                    
                    # Right column: Image (selected size)
                    right_cell = table.rows[0].cells[1]
                    right_cell.width = Inches(image_size)
                    
                    # Set cell width property for stability
                    right_dxa = int(image_size * 1000)  # Convert inches to DXA
                    tcW = parse_xml(f'<w:tcW w:w="{right_dxa}" w:type="dxa" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    right_cell._element.get_or_add_tcPr().append(tcW)
                    
                    # Add vertical alignment
                    vAlign = parse_xml(r'<w:vAlign w:val="center" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    right_cell._element.get_or_add_tcPr().append(vAlign)
                    
                    img_para = right_cell.paragraphs[0]
                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center for better look
                    img_run = img_para.add_run()
                    
                    # Add picture with selected size
                    try:
                        # Use selected image size
                        img_run.add_picture(img_name, width=Inches(image_size))
                    except Exception as img_err:
                        raise Exception(f"Failed to insert image: {str(img_err)}")
                    
                    # Spacing after table
                    spacing_para = doc.add_paragraph()
                    table._tbl.addnext(spacing_para._p)
                    spacing_para.paragraph_format.space_before = Pt(6)
                    spacing_para.paragraph_format.space_after = Pt(18)
                    
                    successful += 1
                    percentage = (image_size / 7.5) * 100
                    print(f"✓ Para {task['para_index']}: Image {image_size}\" ({percentage:.0f}% of page)")
                    
                    # Track temp file for cleanup AFTER document is saved
                    if os.path.exists(img_name):
                        temp_files_to_delete.append(img_name)
                else:
                    print(f"⚠ Para index {task['para_index']} out of range (doc has {len(doc.paragraphs)} paragraphs)")
                    
            except Exception as e:
                failed += 1
                print(f"✗ Failed para {task['para_index']}: {str(e)[:120]}")
                plt.close('all')
                # Failed images can be deleted immediately
                if os.path.exists(img_name):
                    os.remove(img_name)

        # CRITICAL: Save document BEFORE deleting temp image files
        save_path = os.path.join(output_dir, 
                                 os.path.basename(doc_path).replace(".docx", "_Professional.docx"))
        doc.save(save_path)
        
        # NOW it's safe to delete temp files
        print(f"\n🧹 Cleaning up {len(temp_files_to_delete)} temporary image files...")
        for temp_file in temp_files_to_delete:
            if os.path.exists(temp_file):
                os.remove(temp_file)
                print(f"  ✓ Deleted {temp_file}")
        
        percentage = (image_size / 7.5) * 100
        summary = f"""Successfully processed - ULTRA-HD PROFESSIONAL E-BOOK
Saved: {save_path}
✓ Images: {successful} (Size: {image_size}\" - {percentage:.0f}% of page, 600 DPI)
✗ Failed: {failed}
Layout: Dynamic columns, pure black equations, full descriptions"""
        print(summary)
        return summary

class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Professional E-Book Builder")
        self.geometry("750x700")  # Reduced height, will use scrolling
        self.orchestrator = STEMOrchestrator()
        
        # Default image size
        self.selected_image_size = ctk.DoubleVar(value=4.5)
        
        # Create SCROLLABLE frame for all content
        scrollable_frame = ctk.CTkScrollableFrame(self, width=700, height=650)
        scrollable_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        header = ctk.CTkFrame(scrollable_frame, fg_color="#1e3a8a", corner_radius=10)
        header.pack(pady=20, padx=20, fill="x")
        
        ctk.CTkLabel(header, text="Professional E-Book Builder",
                    font=("Arial", 26, "bold"), text_color="white").pack(pady=15)
        
        ctk.CTkLabel(header, text="v12.6 - SCROLLABLE Interface + Adjustable Sizes",
                    font=("Arial", 10), text_color="#93c5fd").pack(pady=(0,15))
        
        # Image Size Selector Frame
        size_frame = ctk.CTkFrame(scrollable_frame, fg_color="#e0f2fe", corner_radius=10)
        size_frame.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkLabel(size_frame, text="📏 Select Image Size:",
                    font=("Arial", 14, "bold"), text_color="#1e3a8a").pack(pady=(10,5))
        
        # Size options in a grid
        sizes_grid = ctk.CTkFrame(size_frame, fg_color="transparent")
        sizes_grid.pack(pady=10, padx=20)
        
        # Create radio buttons for each size
        sizes = [4.0, 4.25, 4.5, 4.75, 5.0, 5.25, 5.5, 5.75, 6.0]
        
        for i, size in enumerate(sizes):
            row = i // 3
            col = i % 3
            
            rb = ctk.CTkRadioButton(
                sizes_grid,
                text=f"{size}\"",
                variable=self.selected_image_size,
                value=size,
                font=("Arial", 12, "bold"),
                fg_color="#2563eb",
                hover_color="#1d4ed8"
            )
            rb.grid(row=row, column=col, padx=15, pady=5, sticky="w")
        
        # Show current selection
        self.size_label = ctk.CTkLabel(
            size_frame, 
            text=f"Current: 4.5 inches (Default)",
            font=("Arial", 11, "bold"),
            text_color="#059669"
        )
        self.size_label.pack(pady=(5,10))
        
        # Update label when selection changes
        self.selected_image_size.trace('w', self.update_size_label)
        
        features = ctk.CTkFrame(scrollable_frame, fg_color="#f0f9ff", corner_radius=10)
        features.pack(pady=15, padx=20, fill="x")
        
        ctk.CTkLabel(features, text="✨ PROFESSIONAL E-BOOK FEATURES",
                    font=("Arial", 13, "bold"), text_color="#1e3a8a").pack(pady=(10,5))
        
        for txt in [
            "📏 Adjustable image sizes: 4.0\" to 6.0\" (0.25\" steps)",
            "📐 Automatic layout adjustment for each size",
            "📝 FULL descriptions - no truncation",
            "📊 Pure BLACK equations (#000000) - maximum contrast",
            "👁️ ULTRA-HD: 600 DPI, razor-sharp quality",
            "📘 Detailed 4-6 line explanations",
            "🔒 100% stable - images never disappear"
        ]:
            ctk.CTkLabel(features, text=txt, font=("Arial", 10),
                        text_color="#374151").pack(pady=1)
        
        ctk.CTkLabel(scrollable_frame, text="Quality: ULTRA-HD 600 DPI | Customizable Sizes | STABLE Layout",
                    font=("Courier", 9), text_color="gray").pack(pady=8)
        
        # === IMPORTANT: FILE UPLOAD BUTTONS - Now inside scrollable frame ===
        ctk.CTkLabel(scrollable_frame, text="👇 SCROLL DOWN TO SEE BUTTONS 👇",
                    font=("Arial", 12, "bold"), text_color="#dc2626").pack(pady=10)
        
        ctk.CTkButton(scrollable_frame, text="📄 Select Single File", command=self.run_file,
                     height=60, width=360, font=("Arial", 16, "bold"),
                     fg_color="#2563eb", hover_color="#1d4ed8").pack(pady=10)
        
        ctk.CTkButton(scrollable_frame, text="📁 Process Entire Folder", command=self.run_folder,
                     height=60, width=360, font=("Arial", 16, "bold"),
                     fg_color="#059669", hover_color="#047857").pack(pady=10)
        
        status_frame = ctk.CTkFrame(scrollable_frame, fg_color="#ecfdf5", corner_radius=8)
        status_frame.pack(pady=10, padx=20, fill="x")
        
        self.status = ctk.CTkLabel(status_frame, text="● Ready - Select Image Size Above, Then Scroll to Buttons",
                                   text_color="#10b981", font=("Arial", 14, "bold"))
        self.status.pack(pady=10)
        
        ctk.CTkLabel(scrollable_frame, text="📊 Console Output:",
                    font=("Arial", 12, "bold")).pack(pady=(10,0), padx=20, anchor="w")
        
        self.console = ctk.CTkTextbox(scrollable_frame, height=150, font=("Courier", 10),
                                      fg_color="#0f172a", text_color="#e2e8f0")
        self.console.pack(pady=5, padx=20, fill="both", expand=True)

    def update_size_label(self, *args):
        size = self.selected_image_size.get()
        # Calculate page percentage (assuming 7.5" page width)
        percentage = (size / 7.5) * 100
        self.size_label.configure(text=f"Current: {size} inches ({percentage:.0f}% of page width)")
        self.status.configure(text=f"● Ready - Image Size: {size}\"", text_color="#10b981")

    def log(self, msg):
        t = datetime.now().strftime("%H:%M:%S")
        self.console.insert("end", f"[{t}] {msg}\n")
        self.console.see("end")
        self.update()

    def run_file(self):
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            # Get selected image size
            img_size = self.selected_image_size.get()
            percentage = (img_size / 7.5) * 100
            
            self.status.configure(text="● Creating E-Book...", text_color="#f59e0b")
            self.console.delete("1.0", "end")
            self.log("═" * 65)
            self.log(f"📄 FILE: {os.path.basename(path)}")
            self.log(f"📏 IMAGE SIZE: {img_size}\" ({percentage:.0f}% of page width)")
            self.log("MODE: Ultra-HD (600 DPI + Black Equations + Full Descriptions)")
            self.log("═" * 65)
            self.update()
            
            result = self.orchestrator.process_file(path, image_size=img_size)
            self.log(f"\n{result}")
            self.log("═" * 65)
            messagebox.showinfo("✓ Complete", result)
            self.status.configure(text=f"● Ready - Image Size: {img_size}\"", text_color="#10b981")

    def run_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            # Get selected image size
            img_size = self.selected_image_size.get()
            percentage = (img_size / 7.5) * 100
            
            self.status.configure(text="● Processing...", text_color="#f59e0b")
            self.console.delete("1.0", "end")
            files = [f for f in os.listdir(folder) if f.endswith(".docx") and not f.startswith("~")]
            self.log("═" * 65)
            self.log(f"📁 FOLDER: {len(files)} documents")
            self.log(f"📏 IMAGE SIZE: {img_size}\" ({percentage:.0f}% of page width)")
            self.log("═" * 65 + "\n")
            
            for i, f in enumerate(files, 1):
                self.log(f"[{i}/{len(files)}] {f}")
                self.orchestrator.process_file(os.path.join(folder, f), image_size=img_size)
                self.log("")
            
            self.log("═" * 65)
            self.log(f"🎉 COMPLETE!")
            self.log("═" * 65)
            messagebox.showinfo("✓ Done", f"Created {len(files)} e-books!")
            self.status.configure(text=f"● Ready - Image Size: {img_size}\"", text_color="#10b981")

if __name__ == "__main__":
    App().mainloop()